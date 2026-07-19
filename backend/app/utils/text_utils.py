"""
文本工具函数
"""
import re
import html
from typing import List, Optional
import unicodedata
import logging

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """清理和规范化文本"""
    if not text:
        return ""
    
    # 移除HTML标签
    text = remove_html_tags(text)
    
    # 规范化unicode
    text = unicodedata.normalize('NFKC', text)
    
    # 移除多余空白字符
    text = re.sub(r'\s+', ' ', text)
    
    # 移除首尾空白字符
    text = text.strip()
    
    return text


def remove_html_tags(text: str) -> str:
    """从文本中移除HTML标签"""
    if not text:
        return ""
    
    # 首先反转义HTML实体
    text = html.unescape(text)
    
    # 移除HTML标签
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', text)
    
    return text


def normalize_text(text: str) -> str:
    """规范化文本用于搜索和比较"""
    if not text:
        return ""
    
    # 转换为小写
    text = text.lower()
    
    # 移除标点符号和特殊字符
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # 移除多余空白字符
    text = re.sub(r'\s+', ' ', text)
    
    # Strip
    text = text.strip()
    
    return text


def truncate_text(text: str, max_length: int, suffix: str = "...") -> str:
    """将文本截断到最大长度并添加后缀"""
    if not text or len(text) <= max_length:
        return text
    
    # 考虑后缀长度
    truncate_length = max_length - len(suffix)
    
    if truncate_length <= 0:
        return suffix[:max_length]
    
    # 尝试在单词边界处截断
    truncated = text[:truncate_length]
    last_space = truncated.rfind(' ')
    
    if last_space > truncate_length * 0.8:  # 如果我们能保留80%的文本
        truncated = truncated[:last_space]
    
    return truncated + suffix


def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """从文本中提取关键词（简单实现）"""
    if not text:
        return []
    
    # 规范化文本
    normalized = normalize_text(text)
    
    # 分割成单词
    words = normalized.split()
    
    # 过滤常见停用词（简单列表）
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have',
        'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
        'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those',
        'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her',
        'us', 'them', 'my', 'your', 'his', 'our', 'their'
    }
    
    # 过滤单词
    keywords = []
    for word in words:
        if (len(word) > 2 and 
            word not in stop_words and 
            word.isalpha()):
            keywords.append(word)
    
    # 统计频率并获取最常见的词
    word_freq = {}
    for word in keywords:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    # 按频率排序并返回关键词
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    
    return [word for word, freq in sorted_words[:max_keywords]]


def split_text_into_chunks(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 100
) -> List[str]:
    """将文本分割成重叠的块"""
    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # 如果不是最后一个块，尝试在句子边界处结束
        if end < len(text):
            # 查找句子结尾
            sentence_end = text.rfind('.', start, end)
            if sentence_end == -1:
                sentence_end = text.rfind('!', start, end)
            if sentence_end == -1:
                sentence_end = text.rfind('?', start, end)

            # 如果找到句子结尾，使用它
            if sentence_end > start + chunk_size * 0.5:
                end = sentence_end + 1
            else:
                # 否则，尝试在单词边界处结束
                space_pos = text.rfind(' ', start, end)
                if space_pos > start + chunk_size * 0.5:
                    end = space_pos

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # 移动起始位置（带重叠）
        start = end - overlap

        # 确保不会倒退
        if start <= chunks[-1] if chunks else 0:
            start = end

    return chunks


def split_options(options_str: str) -> List[str]:
    """按分号分隔符分割选项字符串（支持英文和中文分号）"""
    if not options_str:
        return []

    # 按分号分割（英文和中文）并清理每个选项
    options = [opt.strip() for opt in re.split(r'[;；]', options_str) if opt.strip()]
    return options


def calculate_text_similarity(text1: str, text2: str) -> float:
    """基于共同词汇计算简单的文本相似度"""
    if not text1 or not text2:
        return 0.0
    
    # 规范化文本s
    words1 = set(normalize_text(text1).split())
    words2 = set(normalize_text(text2).split())
    
    if not words1 or not words2:
        return 0.0
    
    # 计算Jaccard相似度
    intersection = len(words1.intersection(words2))
    union = len(words1.union(words2))
    
    return intersection / union if union > 0 else 0.0


def format_file_size(size_bytes: int) -> str:
    """以人类可读的格式格式化文件大小"""
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    size = float(size_bytes)
    
    while size >= 1024.0 and i < len(size_names) - 1:
        size /= 1024.0
        i += 1
    
    return f"{size:.1f} {size_names[i]}"


def escape_markdown(text: str) -> str:
    """转义Markdown特殊字符"""
    if not text:
        return ""
    
    # Markdown中需要转义的字符
    escape_chars = ['\\', '`', '*', '_', '{', '}', '[', ']', '(', ')', '#', '+', '-', '.', '!']
    
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    
    return text


def generate_slug(text: str, max_length: int = 50) -> str:
    """从文本生成URL友好的slug"""
    if not text:
        return ""
    
    # 转换为小写 and normalize
    slug = text.lower()
    slug = unicodedata.normalize('NFKD', slug)
    
    # 移除非ASCII字符
    slug = slug.encode('ascii', 'ignore').decode('ascii')
    
    # 用连字符替换空格和特殊字符
    slug = re.sub(r'[^a-z0-9]+', '-', slug)
    
    # 移除首尾连字符
    slug = slug.strip('-')
    
    # 如果太长则截断
    if len(slug) > max_length:
        slug = slug[:max_length].rstrip('-')
    
    return slug or "untitled"


def mask_sensitive_data(text: str, mask_char: str = "*") -> str:
    """屏蔽敏感数据，如电子邮件、电话号码等"""
    if not text:
        return ""
    
    # 屏蔽电子邮件地址
    text = re.sub(
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        lambda m: m.group(0)[:2] + mask_char * (len(m.group(0)) - 4) + m.group(0)[-2:],
        text
    )
    
    # 屏蔽电话号码（简单模式）
    text = re.sub(
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        lambda m: mask_char * len(m.group(0)),
        text
    )
    
    return text


async def extract_text_content(file_path: str, mime_type: str) -> str:
    """使用健壮的处理程序从文件路径提取文本内容。

    这镜像了增强文档服务的提取逻辑，以便可以在知识库摄取和简历筛选之间重用。
    """
    try:
        if mime_type in ('text/plain', 'text/markdown'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif mime_type == 'application/pdf':
            text = ""
            try:
                import PyPDF2  # type: ignore
            except Exception as e:
                logger.error(f"PyPDF2不可用于PDF提取: {e}")
                return ""

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text += page_text + "\n"
            return text.strip()

        elif mime_type in ('application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Prefer docx2txt for robust extraction
                try:
                    import docx2txt  # type: ignore
                    extracted = docx2txt.process(file_path) or ""
                    if extracted.strip():
                        logger.info("使用docx2txt提取DOCX内容")
                        return extracted.strip()
                    else:
                        logger.info("docx2txt返回空；回退到python-docx")
                except Exception as e:
                    logger.info(f"docx2txt不可用或失败({e})；回退到python-docx")

                # Fallback to python-docx
                try:
                    from docx import Document as DocxDocument  # type: ignore
                    doc = DocxDocument(file_path)
                    parts = []
                    for p in doc.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text and cell.text.strip():
                                    parts.append(cell.text)
                    return "\n".join(parts).strip()
                except Exception as e:
                    logger.error(f"python-docx提取DOCX内容失败: {e}")
                    return ""

            # .doc旧格式：此处不支持，需要外部工具
            logger.warning(f"Unsupported .doc format for direct extraction: {mime_type}")
            return ""

        else:
            logger.warning(f"不支持的文件类型: {mime_type}")
            return ""

    except Exception as e:
        logger.error(f"从{file_path}提取文本时出错: {e}")
        return ""